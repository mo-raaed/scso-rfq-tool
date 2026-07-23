"""Word (.docx) file extractor."""

import os
import re
from typing import List

from models.rfq_data import RFQData, RFQItem
from extractors.base import BaseExtractor


class DOCXExtractor(BaseExtractor):
    """Extracts RFQ data from Word documents."""

    @staticmethod
    def can_handle(file_path: str) -> bool:
        return file_path.lower().endswith(".docx")

    def extract(self, file_path: str) -> RFQData:
        import docx

        data = RFQData(source_file_path=file_path)

        try:
            doc = docx.Document(file_path)
        except Exception as e:
            data.raw_extracted_text = f"[ERROR] Could not open DOCX: {e}"
            return data

        # Extract all paragraph text
        para_texts = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n".join(para_texts)
        data.raw_extracted_text = full_text

        # Extract table data if present
        items: List[RFQItem] = []
        for table in doc.tables:
            for ri, row in enumerate(table.rows):
                if ri == 0:
                    continue  # Skip header row
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) >= 2:
                    desc = cells[1] if len(cells) > 1 else ""
                    qty = cells[2] if len(cells) > 2 else ""
                    if desc or qty:
                        items.append(
                            RFQItem(
                                index=ri,
                                description=desc,
                                quantity=qty,
                            )
                        )

        data.items = items

        # Try to extract metadata from paragraphs
        for text in para_texts:
            stripped = text.strip()
            if stripped.startswith("Tender Subject:"):
                data.tender_subject = stripped.replace("Tender Subject:", "").strip()
            elif stripped.startswith("Closing Date:"):
                data.closing_date = stripped.replace("Closing Date:", "").strip()
            elif stripped.startswith("End") and "User:" in stripped:
                data.end_user = re.sub(r"End[\s\-]*User\s*:", "", stripped).strip()
                data.include_end_user = bool(data.end_user)

        # Fallback subject from filename
        if not data.tender_subject:
            data.tender_subject = os.path.splitext(os.path.basename(file_path))[0]

        return data
