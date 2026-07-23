"""RFQ Builder — populates the SCSO RFQ template and generates output files."""

import os
import re
import shutil
from typing import List, Optional, Callable

import docx
from docx.shared import Pt

from models.rfq_data import RFQData, RFQItem


class RFQBuilder:
    """Builds the final RFQ document(s) from an RFQData object.

    Operations:
      1. Copy template to output location
      2. Populate header fields (Ref #, Subject, Date, etc.)
      3. Remove/keep optional sections based on toggle flags
      4. Clear placeholder table rows and insert real items
      5. Optionally generate manufacturer-split RFQs
      6. Convert to PDF
    """

    def __init__(self, template_path: str, log_callback: Optional[Callable] = None):
        self.template_path = template_path
        self.log = log_callback or print

    # ==================================================================
    # Public API
    # ==================================================================

    def build(self, data: RFQData) -> str:
        """Build the RFQ document and return the output .docx path."""
        output_dir = data.output_folder
        os.makedirs(output_dir, exist_ok=True)

        # Determine output filename
        filename = data.output_filename.strip()
        if not filename:
            filename = f"Request #{data.ref_number} {data.tender_subject}"
        # Sanitize
        filename = re.sub(r'[<>:"/\\|?*]', "_", filename)

        docx_path = os.path.join(output_dir, filename + ".docx")

        # Copy template
        shutil.copy2(self.template_path, docx_path)
        self.log(f"Template copied to: {os.path.basename(docx_path)}")

        # Open and populate
        doc = docx.Document(docx_path)
        self._populate_header(doc, data)
        self._process_optional_sections(doc, data)
        self._update_packing_section(doc, data)
        self._insert_custom_notes(doc, data)
        self._update_engineer_name(doc, data)
        self._populate_table(doc, data)
        doc.save(docx_path)
        self.log(f"Document populated and saved: {os.path.basename(docx_path)}")

        # Convert to PDF
        pdf_path = self._convert_to_pdf(docx_path)
        if pdf_path:
            self.log(f"PDF generated: {os.path.basename(pdf_path)}")

        # Manufacturer-split RFQs
        if data.generate_split_rfqs and data.items:
            self._generate_split_rfqs(docx_path, data)

        return docx_path

    # ==================================================================
    # Header population
    # ==================================================================

    def _populate_header(self, doc: docx.Document, data: RFQData) -> None:
        """Set all header paragraph fields."""
        for p in doc.paragraphs:
            text = p.text.strip()

            # --- Quotation Request: Ref. # ---
            if text.startswith("Quotation Request:"):
                self._set_paragraph_text(
                    p, f"Quotation Request:\t\tRef. # {data.ref_number}"
                )

            # --- Tender Subject: ---
            elif text.startswith("Tender Subject:"):
                subject = data.tender_subject
                if data.tender_number:
                    subject += f" (RFQ: {data.tender_number})"
                self._set_paragraph_text(
                    p, f"Tender Subject:\t\t{subject}"
                )

            # --- Delivery price ---
            elif text.startswith("Delivery price"):
                self._set_paragraph_text(
                    p, f"Delivery price {data.delivery_terms}"
                )

            # --- Closing Date ---
            elif text.startswith("Closing Date:"):
                self._set_paragraph_text(
                    p, f"Closing Date: {data.closing_date}"
                )

            # --- End-User ---
            elif "End" in text and "User:" in text and "Address" not in text:
                if data.include_end_user and data.end_user:
                    # Use non-breaking hyphen for consistency
                    self._set_paragraph_text(p, f"End-User: {data.end_user}")
                else:
                    self._remove_paragraph(p)

            # --- End-User Address ---
            elif "End" in text and "User" in text and "Address" in text:
                if data.include_end_user_address and data.end_user_address:
                    self._set_paragraph_text(
                        p, f"End-User Address: {data.end_user_address}"
                    )
                else:
                    self._remove_paragraph(p)

            # --- Tender Number ---
            elif text.startswith("Tender Number:"):
                if data.include_tender_number and data.tender_number:
                    self._set_paragraph_text(
                        p, f"Tender Number: {data.tender_number}"
                    )
                else:
                    self._remove_paragraph(p)

            # --- Project Name ---
            elif text.startswith("Project Name:"):
                if data.include_project_name and data.project_name:
                    self._set_paragraph_text(
                        p, f"Project Name: {data.project_name}"
                    )
                else:
                    self._remove_paragraph(p)

            # --- Certificate of Origin requirement line ---
            elif text.startswith("Certificate of Origin is required"):
                if data.include_coo_requirement_line:
                    self._set_paragraph_text(p, data.coo_requirement_text)
                else:
                    self._remove_paragraph(p)

            # --- 3rd party inspection requirement line ---
            elif text.startswith("3rd party inspection is required"):
                if data.include_inspection_requirement_line:
                    self._set_paragraph_text(p, data.inspection_requirement_text)
                else:
                    self._remove_paragraph(p)

    # ==================================================================
    # Section Processing
    # ==================================================================

    def _process_optional_sections(
        self, doc: docx.Document, data: RFQData
    ) -> None:
        """Process optional sections: either remove them or replace them with user-customized text."""
        # 1. Body sections: either replace first paragraph and delete rest, or delete all.
        self._process_body_section(
            doc, 
            data.include_inspection_section, 
            "3rd Party Inspection:", 
            "Certificate of Origin (COO):", 
            data.inspection_section_text
        )
        
        self._process_body_section(
            doc, 
            data.include_coo_section, 
            "Certificate of Origin (COO):", 
            "General Note:", 
            data.coo_section_text
        )
        
        self._process_body_section(
            doc, 
            data.include_general_note, 
            "General Note:", 
            "Thank You and Best Regards", 
            data.general_note_text
        )

        # 2. Country of origin section
        self._process_country_of_origin(doc, data)

    def _process_body_section(
        self, doc: docx.Document, include: bool, start_marker: str, end_marker: str, custom_text: str
    ) -> None:
        """Helper to replace a section header and delete subsequent details, or delete entire section."""
        paragraphs_to_remove = []
        first_p = None
        in_section = False
        
        for p in doc.paragraphs:
            text = p.text.strip()
            if text.startswith(start_marker):
                in_section = True
                first_p = p
                continue
            if in_section:
                if text.startswith(end_marker):
                    in_section = False
                    break
                paragraphs_to_remove.append(p)
                
        if include and first_p:
            self._replace_paragraph_with_multiline(first_p, custom_text)
            for p in paragraphs_to_remove:
                self._remove_paragraph(p)
        elif not include:
            if first_p:
                self._remove_paragraph(first_p)
            for p in paragraphs_to_remove:
                self._remove_paragraph(p)

    def _replace_paragraph_with_multiline(self, paragraph, text: str) -> None:
        """Replace a paragraph's content with multi-line text, clearing highlight and justification."""
        lines = [ln.strip() for ln in text.splitlines()]
        if not lines:
            self._remove_paragraph(paragraph)
            return

        # 1. Update the first paragraph
        paragraph.text = "" # Clears all runs and formatting/highlighting in runs
        paragraph.paragraph_format.alignment = None # Reset justification to left-aligned
        
        # Clear paragraph-level shading if present in XML
        try:
            pPr = paragraph._element.get_or_add_pPr()
            shd = pPr.find(docx.oxml.ns.qn('w:shd'))
            if shd is not None:
                pPr.remove(shd)
        except Exception:
            pass
        
        # Add the first line to the first paragraph
        run = paragraph.add_run(lines[0])
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Cambria"

        # 2. Insert subsequent lines as separate paragraphs
        from docx.oxml import OxmlElement
        from docx.text.paragraph import Paragraph
        
        current_p = paragraph
        for line in lines[1:]:
            new_p_element = OxmlElement('w:p')
            current_p._p.addnext(new_p_element)
            new_p = Paragraph(new_p_element, paragraph._parent)
            
            # Apply layout styles (left aligned, standard spacing)
            new_p.paragraph_format.alignment = None
            new_p.paragraph_format.space_after = Pt(6)
            
            # Add text
            new_run = new_p.add_run(line)
            new_run.font.size = Pt(11)
            new_run.font.name = "Calibri"
            
            # Move reference forward
            current_p = new_p

    def _process_country_of_origin(self, doc: docx.Document, data: RFQData) -> None:
        """Insert or clean up Country of Origin section."""
        coo_p = None
        in_coo_section = False
        coo_paras_to_remove = []
        
        for p in doc.paragraphs:
            text = p.text.strip()
            if text.startswith("Country of Origin:"):
                coo_p = p
                in_coo_section = True
                continue
            if in_coo_section:
                if (text.startswith("3rd Party Inspection:") or 
                    text.startswith("General Note:") or 
                    text.startswith("Thank You and Best Regards") or 
                    text.startswith("Certificate of Origin (COO):")):
                    in_coo_section = False
                    break
                coo_paras_to_remove.append(p)

        if coo_p:
            self._remove_paragraph(coo_p)
            for p in coo_paras_to_remove:
                self._remove_paragraph(p)

        if data.include_country_of_origin and data.country_of_origin_text:
            packing_end_p = None
            for p in doc.paragraphs:
                if p.text.strip().startswith("Please inform us of the estimated weight"):
                    packing_end_p = p
                    break
            
            if packing_end_p:
                parent = packing_end_p._element.getparent()
                idx = list(parent).index(packing_end_p._element)
                
                # Create a new paragraph and add the custom text
                new_p = docx.oxml.OxmlElement("w:p")
                new_r = docx.oxml.OxmlElement("w:r")
                new_t = docx.oxml.OxmlElement("w:t")
                new_t.text = "\n" + data.country_of_origin_text
                new_r.append(new_t)
                new_p.append(new_r)
                parent.insert(idx + 1, new_p)

    # ==================================================================
    # Packing section
    # ==================================================================

    def _update_packing_section(
        self, doc: docx.Document, data: RFQData
    ) -> None:
        """Update the packing section text."""
        for p in doc.paragraphs:
            text = p.text.strip()

            if text == "Packing Requirements":
                self._set_paragraph_text(p, data.packing_header)

            elif text == "As per manufacturer standards.":
                self._set_paragraph_text(p, data.packing_text)

            elif text.startswith(
                "Please inform us of the estimated weight"
            ):
                self._set_paragraph_text(p, data.packing_detail)

    # ==================================================================
    # Custom notes
    # ==================================================================

    def _insert_custom_notes(
        self, doc: docx.Document, data: RFQData
    ) -> None:
        """Insert custom note paragraphs before 'SCOPE of Supply:'."""
        if not data.custom_notes:
            return

        # Find the "SCOPE of Supply:" paragraph
        scope_para = None
        for p in doc.paragraphs:
            if p.text.strip().startswith("SCOPE of Supply"):
                scope_para = p
                break

        if not scope_para:
            return

        # Insert custom notes before SCOPE of Supply
        parent = scope_para._element.getparent()
        scope_idx = list(parent).index(scope_para._element)

        for note_text in reversed(data.custom_notes):
            if not note_text.strip():
                continue
            new_p = docx.oxml.OxmlElement("w:p")
            new_r = docx.oxml.OxmlElement("w:r")
            new_t = docx.oxml.OxmlElement("w:t")
            new_t.text = note_text
            new_r.append(new_t)
            new_p.append(new_r)
            parent.insert(scope_idx, new_p)

    # ==================================================================
    # Engineer name
    # ==================================================================

    def _update_engineer_name(
        self, doc: docx.Document, data: RFQData
    ) -> None:
        """Set the engineer's name at the bottom."""
        for p in doc.paragraphs:
            text = p.text.strip()
            if text.startswith("Eng."):
                self._set_paragraph_text(p, f"Eng. {data.engineer_name}")

    # ==================================================================
    # Table population
    # ==================================================================

    def _populate_table(self, doc: docx.Document, data: RFQData) -> None:
        """Clear placeholder rows and insert real item data."""
        if not doc.tables:
            self.log("[Warning] No table found in template.")
            return

        table = doc.tables[0]

        # Preserve header row formatting
        header_row = table.rows[0] if table.rows else None

        # Clear all rows after header
        while len(table.rows) > 1:
            tr = table.rows[1]._tr
            tr.getparent().remove(tr)

        # Check if we should group by manufacturer
        if data.generate_split_rfqs:
            # For the general doc when splits exist, group by manufacturer
            self._add_grouped_rows(table, data.items)
        else:
            # Add items sequentially
            for idx, item in enumerate(data.items, 1):
                row = table.add_row()
                row.cells[0].text = str(idx).zfill(2)
                row.cells[1].text = item.description
                row.cells[2].text = str(item.quantity)
                self._apply_cell_font(row)

        self.log(f"Table populated with {len(data.items)} item(s).")

    def _add_grouped_rows(
        self, table, items: List[RFQItem]
    ) -> None:
        """Add rows grouped by manufacturer (Petronas-style)."""
        groups = []
        seen = {}

        for item in items:
            mfg = (item.manufacturer or "").strip().lower()
            if mfg and mfg in seen:
                groups[seen[mfg]]["items"].append(item)
            else:
                if mfg:
                    seen[mfg] = len(groups)
                groups.append({
                    "manufacturer": item.manufacturer,
                    "items": [item],
                })

        for idx, group in enumerate(groups, 1):
            row = table.add_row()
            row.cells[0].text = str(idx)

            descs = [it.description for it in group["items"]]
            row.cells[1].text = "\n".join(descs)

            qtys = [str(it.quantity) for it in group["items"]]
            row.cells[2].text = "\n".join(qtys)

            self._apply_cell_font(row)

    # ==================================================================
    # Manufacturer-split RFQs
    # ==================================================================

    def _generate_split_rfqs(
        self, general_docx_path: str, data: RFQData
    ) -> None:
        """Create per-manufacturer RFQ documents."""
        # Group items by manufacturer
        mfg_items = {}
        for item in data.items:
            mfg = (item.manufacturer or "").strip()
            if mfg:
                mfg_items.setdefault(mfg, []).append(item)

        if not mfg_items:
            self.log("No manufacturers found for split RFQs.")
            return

        self.log(
            f"Generating {len(mfg_items)} manufacturer-specific RFQ(s)..."
        )

        for mfg, items in mfg_items.items():
            # Sanitize filename
            mfg_safe = re.sub(r'[<>:"/\\|?*]', "_", mfg)
            mfg_filename = f"{data.tender_subject} - {mfg_safe}.docx"
            mfg_path = os.path.join(data.output_folder, mfg_filename)

            shutil.copy2(general_docx_path, mfg_path)
            doc = docx.Document(mfg_path)

            # Update subject to include manufacturer
            for p in doc.paragraphs:
                if p.text.strip().startswith("Tender Subject:"):
                    subject = data.tender_subject
                    if data.tender_number:
                        subject += f" - {mfg} (RFQ: {data.tender_number})"
                    else:
                        subject += f" - {mfg}"
                    self._set_paragraph_text(
                        p, f"Tender Subject:\t\t{subject}"
                    )
                    break

            # Rebuild table with only this manufacturer's items
            if doc.tables:
                table = doc.tables[0]
                while len(table.rows) > 1:
                    tr = table.rows[1]._tr
                    tr.getparent().remove(tr)

                for idx, item in enumerate(items, 1):
                    row = table.add_row()
                    row.cells[0].text = str(idx)
                    row.cells[1].text = item.description
                    row.cells[2].text = str(item.quantity)
                    self._apply_cell_font(row)

            doc.save(mfg_path)
            self.log(f"  Saved: {mfg_filename}")

            # Convert split RFQ to PDF too
            self._convert_to_pdf(mfg_path)

    # ==================================================================
    # PDF conversion
    # ==================================================================

    def _convert_to_pdf(self, docx_path: str) -> Optional[str]:
        """Convert .docx to .pdf using dynamic win32com dispatch to prevent cache errors under PyInstaller."""
        pdf_path = docx_path.replace(".docx", ".pdf")
        try:
            import win32com.client
            abs_docx = os.path.abspath(docx_path)
            abs_pdf = os.path.abspath(pdf_path)
            
            word = None
            doc = None
            try:
                # Try dynamic dispatch to prevent PyInstaller gen_py cache crashes
                try:
                    word = win32com.client.dynamic.Dispatch("Word.Application")
                except Exception:
                    word = win32com.client.Dispatch("Word.Application")
                
                word.Visible = False
                word.DisplayAlerts = 0  # wdAlertsNone = 0
                
                doc = word.Documents.Open(abs_docx, ConfirmConversions=False, ReadOnly=True)
                doc.SaveAs(abs_pdf, FileFormat=17) # wdFormatPDF = 17
                return pdf_path
            finally:
                if doc is not None:
                    try:
                        doc.Close(0) # wdDoNotSaveChanges = 0
                    except Exception:
                        pass
                if word is not None:
                    try:
                        word.Quit()
                    except Exception:
                        pass
        except Exception as e:
            self.log(
                f"[Warning] PDF conversion failed (is MS Word installed?): {e}"
            )
            return None

    # ==================================================================
    # Helpers
    # ==================================================================

    @staticmethod
    def _set_paragraph_text(p, new_text: str) -> None:
        """Replace paragraph text while preserving first run's formatting."""
        if p.runs:
            p.runs[0].text = new_text
            for r in p.runs[1:]:
                r.text = ""
        else:
            p.text = new_text

    @staticmethod
    def _remove_paragraph(p) -> None:
        """Remove a paragraph element from the document."""
        try:
            parent = p._element.getparent()
            if parent is not None:
                parent.remove(p._element)
        except Exception:
            pass

    @staticmethod
    def _apply_cell_font(row) -> None:
        """Apply consistent font to all cells in a row."""
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Calibri"
